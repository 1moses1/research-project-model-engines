"""
DOCX Text Extraction Service
"""

from docx import Document


class DOCXExtractor:
    """Extracts text from DOCX files"""

    def extract(self, file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            print(f"📄 Extracted {len(text_content)} paragraphs/rows")

            full_text = "\n".join(text_content)

            return full_text

        except Exception as e:
            print(f"⚠️ DOCX extraction error: {str(e)}")
            return f"Error extracting DOCX: {str(e)}"
